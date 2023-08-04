from flask import Flask, request, jsonify
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
import os
import docx
from dotenv import load_dotenv
from docx import Document

load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

app = Flask(__name__)


template = """Your task is to use the Main Input: {main_input}  and the four sample inputs provided which are
Sample Input 1: {sample_input_1}
Sample Input 2: {sample_input_2}
Sample Input 3: {sample_input_3}
Sample Input 4: {sample_input_4}

to analyze the information and propose the best possible input.
Be creative, thoughtful, and strategic in your approach. The quality of your proposed input 
will be based on how well you utilize the given data to craft an effective solution.

Best Input:"""

prompt = PromptTemplate(input_variables=["main_input", "sample_input_1", "sample_input_2", "sample_input_3", "sample_input_4"], template=template)

def extract_paragraphs_from_docx(docx_path):
    paragraphs = []
    doc = docx.Document(docx_path)
    for para in doc.paragraphs:
        text = para.text.strip()
        # Ignore empty paragraphs and paragraphs with only spaces
        if text:
            paragraphs.append(para)
    return paragraphs

def get_paragraph_style(para):
    para_style = para.style.name
    return para_style

def is_heading_style(para):
    heading_styles = ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6']
    return any(style in get_paragraph_style(para) for style in heading_styles)

def get_key_from_paragraph(para):
    first_word = para.text.strip().split()[0]
    if is_heading_style(para):
        return para.text.strip().rstrip('.')
    elif is_heading_style(para.runs[0]):
        return " ".join(para.text.strip().split()[:3]).rstrip('.')
    return None

def create_topics_map(paragraphs):
    topics_map = {}
    current_topic = None
    related_paragraphs = []
    for para in paragraphs:
        key = get_key_from_paragraph(para)

        if key:
            if current_topic:
                current_topic = current_topic.strip().rstrip('.')
                topics_map[current_topic] = related_paragraphs
                related_paragraphs = []

            current_topic = key

        if current_topic:
            related_paragraphs.append(para.text.strip())

    if current_topic is not None:
        current_topic = current_topic.strip().rstrip('.')
        topics_map[current_topic] = related_paragraphs

    return topics_map



llm = ChatOpenAI(model_name="gpt-3.5-turbo", openai_api_key=openai_api_key, temperature=0)
finetune_chain = LLMChain(llm=llm, prompt=prompt)

@app.route('/find_paragraphs', methods=['POST'])
def search_topic_paragraphs():
    request_data = request.json
    topics = request_data.get('topics', [])
    file_path = request_data.get('docx_file_path', '')

    if not file_path:
        return jsonify({"error": "Document file path not provided."}), 400

    # Load the document and create topics_map
    paragraphs = extract_paragraphs_from_docx(file_path)
    topics_map = create_topics_map(paragraphs)

    result = {}
    for topic in topics:
        if topic in topics_map:
            result[topic] = topics_map[topic]
        else:
            result[topic] = None

    return jsonify(result)



@app.route('/get_best_input', methods=['POST'])
def get_best_input():
    data = request.get_json()

    main_input = data.get('main_input', '')
    sample_input_1 = data.get('sample_input_1', '')
    sample_input_2 = data.get('sample_input_2', '')
    sample_input_3 = data.get('sample_input_3', '')
    sample_input_4 = data.get('sample_input_4', '')

    best_input = finetune_chain.predict(main_input=main_input, sample_input_1=sample_input_1,
                               sample_input_2=sample_input_2, sample_input_3=sample_input_3,
                               sample_input_4=sample_input_4)

    return jsonify({'best_input': best_input})

@app.route('/question-answering', methods=['POST'])
def question_answering():
    data = request.get_json()
    doc_path = data.get('doc_path')
    query = data.get('query')

    doc = Document(doc_path)
    raw_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=800,
        chunk_overlap=10,
    )
    texts = text_splitter.split_text(raw_text)

    embeddings = OpenAIEmbeddings()
    docsearch = FAISS.from_texts(texts, embeddings)
    llm = ChatOpenAI(model_name="gpt-3.5-turbo", openai_api_key=openai_api_key, temperature=0)
    chain = load_qa_chain(llm=llm, chain_type="stuff")

    docs = docsearch.similarity_search(query)
    answer = chain.run(input_documents=docs, question=query)

    return jsonify({'answer': answer})

@app.route('/')
def hello_world():
    return 'welcome to flask app'


if __name__ == '__main__':
    app.run(debug=True,port=8000)
